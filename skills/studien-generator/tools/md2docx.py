#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md2docx.py — Markdown -> DOCX (JANS-Layout) + optional PDF.

Setzt die Dauerregel um: zu jedem inhaltlichen .md-Erzeugnis eine DOCX- und PDF-Fassung
mit gleichem Namensstamm daneben ablegen (Rule auto-verbesserungen 260603).

Layout: Cambria 11pt, A4, 20 mm Raender, schwarz, dezente Tabellenlinien
(angelehnt an dokument-layout-standard.md / jans_docx.py).

Unterstuetzt: # / ## / ### Ueberschriften, Pipe-Tabellen, Aufzaehlungen (-, *),
nummerierte Listen, Blockzitate (>), Code-Bloecke (```), **fett**, `code`, --- Trennlinie.

⚠ OFFEN seit 14.08.2026 (Synergie-Lauf 07, SYN-27): Die Rule dokument-layout-standard.md
verbietet in AUSGEHENDEN Dokumenten seit dem 14.08.2026 Aufzaehlungs-Bullets (weder «•»
noch «-») und verlangt stattdessen nummerierte Positionen im Fliesstext (Themenblock 01,
02 …; Positionen 01.1, 01.2 … mit fetter Nummer und fettem Stichwort). Dieser Konverter
setzt eine «- »-Zeile weiterhin als Word-Bullet (style="List Bullet", siehe unten) und ist
damit NICHT auf dem Stand der Rule. Bis zur Umstellung gilt: fuer ausgehende Dokumente die
Positionen bereits in der MD-Quelle nummeriert schreiben, nicht als Bullet-Liste — der
Konverter heilt das nicht. Interne Arbeitsdokumente duerfen Bullets behalten.

Zeilenumbrueche (seit 30.07.2026): EINGERUECKTE Fortsetzungszeilen werden an ihren Block
angehaengt (siehe _unwrap). Damit werden umbrochene Listenpunkte wieder ein Absatz und ein
ueber den Umbruch laufendes **fett** wird erkannt. NICHT eingerueckte Folgezeilen bleiben
bewusst eigene Absaetze, damit Adress-/Signaturbloecke nicht zu einer Zeile verschmelzen.

Aufruf:
  python3 md2docx.py <datei.md> [--pdf]        # docx neben md; --pdf erzeugt auch PDF
  python3 md2docx.py <datei.md> -o <ziel.docx> [--pdf]

PDF-Export via LibreOffice:  soffice --headless --convert-to pdf <datei.docx>
"""
import argparse, os, re, subprocess, sys
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Cambria"
SIZE = 11
BLACK = RGBColor(0, 0, 0)

SOFFICE_CANDIDATES = [
    "/opt/homebrew/bin/soffice", "/usr/local/bin/soffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice", "soffice",
]


def _base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    st.font.color.rgb = BLACK
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for s in doc.sections:
        s.page_height, s.page_width = Mm(297), Mm(210)
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Mm(20)
    for name, sz in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        try:
            h = doc.styles[name]
            h.font.name = FONT
            h.font.size = Pt(sz)
            h.font.color.rgb = BLACK
            h.font.bold = True
        except KeyError:
            pass


def _runs(p, text):
    """Inline **fett** und `code` aufloesen."""
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            p.add_run(tok)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def _table(doc, rows):
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(cols):
            txt = row[ci] if ci < len(row) else ""
            para = cells[ci].paragraphs[0]
            _runs(para, txt.strip())
            for run in para.runs:
                run.font.name = FONT; run.font.size = Pt(10)
                if ri == 0:
                    run.bold = True
            if ri == 0:
                _shade(cells[ci], "F2F2F2")


BLOCK_START = re.compile(r"^\s*(#{1,6}\s|[-*]\s|\d+\.\s|>|\||```|---+\s*$)")


def _unwrap(lines):
    """Eingerueckte Fortsetzungszeilen an ihren Block anhaengen.

    Warum (30.07.2026, vollgas-radar): jede physische Zeile wurde bisher zu einem
    eigenen Absatz. Ein ueber den Umbruch laufendes **fett** war damit nie
    erkennbar (die Inline-Regex arbeitet zeilenweise) und Listenpunkte zerfielen
    in Bruchstuecke — der Normalfall in jeder raw-Datei des Wissens-Layers, die
    auf ~95 Zeichen umbrochen ist.

    BEWUSST ENG: zusammengefuehrt wird NUR, wenn die Folgezeile eingerueckt ist.
    Nicht eingerueckte Folgezeilen bleiben eigene Absaetze wie bisher — sonst
    wuerden Adress- und Signaturbloecke (Briefkopf, Absenderblock) zu einer
    einzigen Zeile verschmelzen. Markdown-konform waere das Zusammenfuehren,
    JANS-konform ist es nicht.
    """
    out, buf, in_code = [], None, False

    def flush():
        nonlocal buf
        if buf is not None:
            out.append(buf)
            buf = None

    for ln in lines:
        if ln.strip().startswith("```"):
            flush(); out.append(ln); in_code = not in_code; continue
        if in_code:
            out.append(ln); continue
        if not ln.strip():
            flush(); out.append(""); continue
        indented = ln[:1].isspace()
        if BLOCK_START.match(ln) or not indented or buf is None:
            flush(); buf = ln.rstrip(); continue
        prev = buf.rstrip()
        # Pfad oder Wort lief ueber den Umbruch: ohne Leerschlag zusammenziehen.
        buf = prev + ("" if prev.endswith(("/", "-")) else " ") + ln.strip()
    flush()
    return out


def convert(md_path, out_path=None):
    with open(md_path, encoding="utf-8") as f:
        lines = _unwrap(f.read().splitlines())
    doc = Document(); _base(doc)
    i, n = 0, len(lines)
    in_code = False; code_buf = []
    while i < n:
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                r.font.name = "Consolas"; r.font.size = Pt(9)
                code_buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue
        # Tabelle
        if line.lstrip().startswith("|") and "|" in line[1:]:
            block = []
            while i < n and lines[i].lstrip().startswith("|"):
                block.append(lines[i]); i += 1
            rows = []
            for r in block:
                if re.match(r"^\s*\|?[\s:|-]+\|?\s*$", r):  # ---|--- Trennzeile
                    continue
                cells = [c for c in r.strip().strip("|").split("|")]
                rows.append(cells)
            if rows:
                _table(doc, rows)
            continue
        # Ueberschriften
        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            doc.add_heading(m.group(2).strip(), level=len(m.group(1))); i += 1; continue
        # Trennlinie
        if re.match(r"^\s*---+\s*$", line):
            i += 1; continue
        # Blockzitat
        if line.lstrip().startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Mm(6)
            r = p.add_run(line.lstrip().lstrip(">").strip()); r.italic = True
            i += 1; continue
        # Aufzaehlung
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Mm(6 + 4 * (len(m.group(1)) // 2))
            _runs(p, m.group(2)); i += 1; continue
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number"); _runs(p, m.group(1)); i += 1; continue
        # Leerzeile / Normal
        if not line.strip():
            i += 1; continue
        p = doc.add_paragraph(); _runs(p, line); i += 1
    out = out_path or os.path.splitext(md_path)[0] + ".docx"
    doc.save(out)
    return out


def to_pdf(docx_path):
    soffice = next((c for c in SOFFICE_CANDIDATES
                    if os.path.exists(c) or c == "soffice"), None)
    outdir = os.path.dirname(os.path.abspath(docx_path))
    subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                    "--outdir", outdir, docx_path], check=True,
                   stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return os.path.splitext(docx_path)[0] + ".pdf"


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("md")
    ap.add_argument("-o", "--out")
    ap.add_argument("--pdf", action="store_true")
    a = ap.parse_args()
    docx = convert(a.md, a.out)
    print("DOCX:", docx)
    if a.pdf:
        print("PDF :", to_pdf(docx))


if __name__ == "__main__":
    main()
