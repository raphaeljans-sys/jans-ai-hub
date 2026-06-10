#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_grobkosten_onepager.py — JANS Grobkosten-Onepager (eine A4-Seite, praesentierwuerdig)

Skill `grobkosten-onepager`.

Erzeugt aus einer einfachen Konfiguration (Volumen x Kennwert) eine einseitige,
auf den ersten Blick lesbare Grobkostenschaetzung im JANS-Layout (Cambria, schwarz,
rahmenlos) mit grossen Kosten-Kennzahlen je Variante.

Methode bewusst simpel: Erstellungskosten ~= Gebaeudevolumen x Kennwert (BKP 1-5).

Aufruf:
  python build_grobkosten_onepager.py <config.json> <ausgabe.docx>

config.json (Beispiel siehe skills/grobkosten-onepager/SKILL.md):
  {
    "titel": "...", "zeile2": "...", "datum": "070626",
    "summary": "...",                         # optionaler Headline-Satz
    "kennwert": 950, "band_low": 850, "band_high": 1050,
    "einheit": "m3", "bkp": "BKP 1-5",
    "kennwert_quelle": "JANS Erfahrungswert ...",
    "varianten": [
       {"name":"A Steildach","kurz":"...","volumen":1600,"zuschlag_pct":0,"zuschlag_label":""},
       ...
    ],
    "herleitung": ["...","..."],
    "vorbehalte": ["...","..."],
    "footer": "..."
  }

Abhaengigkeit: python-docx; PDF-Export via scripts/docx2pdf.sh.
Kennwerte-Quelle (Lernmodell): wissen/grobkosten/wiki/kennwerte.md
"""
import sys, os, json

# jans_docx (Layout-Bausteine) aus dem Schwester-Skill ausschreibung beziehen
_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.join(_HERE, "..", "..", "ausschreibung", "tools"))
from jans_docx import base_doc, para, BLACK, FONT, JANS_NAME, JANS_STREET, JANS_CITY  # noqa: E402
from docx.shared import Pt, Mm, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402

CENTER = WD_ALIGN_PARAGRAPH.CENTER


def _mio(x):
    return f"CHF {x/1e6:.1f} Mio".replace(".0 Mio", " Mio")


def _band(low, high):
    return f"{low/1e6:.1f}–{high/1e6:.1f} Mio"


def _no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'), 'none'); b.append(e)
    tblPr.append(b)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)


def _set_grid(tbl, widths_mm):
    grid = tbl._tbl.tblGrid
    for col, w in zip(grid.findall(qn('w:gridCol')), widths_mm):
        col.set(qn('w:w'), str(int(w * 56.6929)))


def _cellpar(cell, text, size, bold=False, after=2, align=CENTER):
    p = cell.paragraphs[0] if not cell.paragraphs[0].runs else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = BLACK
    return p


def _top_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    top = OxmlElement('w:top'); top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6'); top.set(qn('w:color'), '000000')
    borders.append(top); tcPr.append(borders)


def build(cfg, out_path):
    d = base_doc()
    # Kopf
    para(d, cfg["titel"], size=17, bold=True, after=2)
    para(d, cfg.get("zeile2", ""), size=10, after=1)
    para(d, f"{JANS_NAME}  ·  {JANS_STREET}  ·  {JANS_CITY}  ·  Stand {cfg.get('datum','')}",
         size=9, after=6)

    # Headline-Satz (auf den ersten Blick)
    if cfg.get("summary"):
        para(d, cfg["summary"], size=12, bold=True, after=2)

    # Methode in einer Zeile
    para(d, f"Methode: Erstellungskosten ≈ Gebäudevolumen × Kennwert ({cfg.get('bkp','BKP 1-5')}). "
            f"Kennwert {int(cfg['kennwert'])} CHF/{cfg.get('einheit','m3')} "
            f"(Bandbreite {int(cfg['band_low'])}–{int(cfg['band_high'])}).",
         size=10, after=8)

    # Varianten-Karten (3-spaltig, grosse Kennzahl)
    vs = cfg["varianten"]
    n = len(vs)
    total_w = 170.0
    widths = [total_w / n] * n
    tbl = d.add_table(rows=1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False; tbl.allow_autofit = False
    _no_borders(tbl); _set_grid(tbl, widths)
    for i, v in enumerate(vs):
        vol = float(v["volumen"]); z = float(v.get("zuschlag_pct", 0) or 0) / 100.0
        kosten = vol * cfg["kennwert"] * (1 + z)
        low = vol * cfg["band_low"] * (1 + z)
        high = vol * cfg["band_high"] * (1 + z)
        cell = tbl.rows[0].cells[i]
        _cellpar(cell, v["name"], 13, bold=True, after=3)
        if v.get("kurz"):
            _cellpar(cell, v["kurz"], 9, after=4)
        _cellpar(cell, f"Volumen  ~{int(round(vol)):,} {cfg.get('einheit','m3')}".replace(",", "'"), 10, after=1)
        zl = f"  (+{int(v['zuschlag_pct'])}% {v.get('zuschlag_label','')})" if z else ""
        _cellpar(cell, f"× {int(cfg['kennwert'])} CHF/{cfg.get('einheit','m3')}{zl}", 10, after=4)
        _cellpar(cell, _mio(kosten), 22, bold=True, after=1)
        _cellpar(cell, f"Bandbreite {_band(low, high)}", 9, after=2)

    para(d, "", after=4)

    # Herleitung / Kennwert-Box
    if cfg.get("herleitung"):
        para(d, "Herleitung Kennwert", size=11, bold=True, before=6, after=2)
        for h in cfg["herleitung"]:
            para(d, "– " + h, size=10, after=1)
    if cfg.get("kennwert_quelle"):
        para(d, "Kennwert-Quelle: " + cfg["kennwert_quelle"], size=9, before=2, after=2)

    # Vorbehalte
    if cfg.get("vorbehalte"):
        para(d, "Annahmen & Vorbehalte", size=11, bold=True, before=6, after=2)
        for vb in cfg["vorbehalte"]:
            para(d, "– " + vb, size=9, after=1)

    if cfg.get("footer"):
        para(d, cfg["footer"], size=8, before=8, align=CENTER)

    d.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Aufruf: build_grobkosten_onepager.py <config.json> <ausgabe.docx>"); sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        cfg = json.load(f)
    print("OK:", build(cfg, sys.argv[2]))
