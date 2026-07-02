#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
jans_docx.py — Wiederverwendbarer JANS-Dokument-Baukasten (DOCX, JANS-Layout)

JANS AI Hub · Skill `ausschreibung`

Baut LV und Begleitschreiben (Anschreiben) im verbindlichen JANS-Layout
(Cambria 11pt, A4, schwarz, rahmenlos) — inkl. Seitenzahlen, korrekter
Absenderadresse und allen Pflicht-Elementen. Damit greifen alle erarbeiteten
Standards bei jeder Ausschreibung automatisch.

Abhaengigkeit: python-docx. PDF-Export anschliessend via:
  soffice --headless --convert-to pdf <datei>.docx

Verbindliche Standards, die hier fest verdrahtet sind:
- Layout: Cambria 11pt, A4, 20mm Raender, schwarz, keine sichtbaren Rahmen
- Footer mit Seitenzahl "Seite X von Y"
- Absenderadresse: Grubenstrasse 37, 8045 Zuerich (Rule jans-absenderadresse.md)
- MWST 8,1 %
- Beilagen / Ausschreibungsunterlagen: pro Zeile aufgelistet
- Plan-/Dokumentengrundlagen werden konkret benannt
- "Gleichwertig"-Klausel + Bauseits-Hinweis
- Optionaler Besichtigungstermin
"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
FONT = "Cambria"

# JANS Absender (Quelle: rules/jans-absenderadresse.md — bei Aenderung dort UND hier nachfuehren)
JANS_NAME = "Raphael Jans Architekten ETH/SIA"
JANS_STREET = "Grubenstrasse 37"
JANS_CITY = "8045 Zürich"
JANS_MAIL = "rj@raphaeljans.ch"
JANS_WEB = "www.raphaeljans.ch"
MWST = "8,1 %"


def base_doc():
    d = Document()
    s = d.sections[0]
    s.page_height, s.page_width = Mm(297), Mm(210)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Mm(20)
    st = d.styles["Normal"]
    st.font.name = FONT; st.font.size = Pt(11); st.font.color.rgb = BLACK
    st.paragraph_format.line_spacing = 1.4
    st.paragraph_format.space_after = Pt(0)
    return d


def para(d, text="", size=11, bold=False, before=0, after=0, align=None):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = BLACK
    return p


def h1(d, t): para(d, t, size=18, bold=True, after=6)
def h2(d, t): para(d, t, size=12, bold=True, before=14, after=4)


def _no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'), 'none'); borders.append(e)
    tblPr.append(borders)
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)


def _set_grid(tbl, widths_mm):
    grid = tbl._tbl.tblGrid
    for col, w in zip(grid.findall(qn('w:gridCol')), widths_mm):
        col.set(qn('w:w'), str(int(w * 56.6929)))


_ALIGN = {"l": WD_ALIGN_PARAGRAPH.LEFT, "r": WD_ALIGN_PARAGRAPH.RIGHT, "c": WD_ALIGN_PARAGRAPH.CENTER}


def _tight_margins(tbl, lr=0.6):
    """Schmale Zell-Innenraender (Standard-Word: 1.9 mm L/R) -> mehr nutzbare Spaltenbreite,
    damit Einheiten/Zahlen/Header nicht unnoetig umbrechen (Lesbarkeits-Standard)."""
    tblPr = tbl._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for edge, val in (('top', 0), ('bottom', 0), ('left', lr), ('right', lr)):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:w'), str(int(val * 56.6929))); e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tblPr.append(mar)


def _set_cell(cell, text, size=10, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.2
    if align in _ALIGN:
        p.alignment = _ALIGN[align]
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = BLACK


def table(d, rows, widths, header=True, aligns=None):
    """Rahmenlose, fixbreite Tabelle (Summe widths <= 170 mm). header=True -> Zeile 0 fett.

    aligns: optionale Liste je Spalte ('l'/'r'/'c'), z.B. Zahlenspalten rechtsbuendig.
    Spalten so breit waehlen, dass Inhalt (inkl. Header und Einheiten) NICHT umbricht
    — fehlerhafte Umbrueche sind dem Lesbarkeits-Standard abtraeglich.
    """
    tbl = d.add_table(rows=0, cols=len(widths))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False; tbl.allow_autofit = False
    _no_borders(tbl); _set_grid(tbl, widths); _tight_margins(tbl)
    for ri, r in enumerate(rows):
        cells = tbl.add_row().cells
        for i, val in enumerate(r):
            a = aligns[i] if aligns and i < len(aligns) else None
            _set_cell(cells[i], val, size=10, bold=(header and ri == 0), align=a)
            cells[i].width = Mm(widths[i])
    return tbl


def _ft_run(p, text):
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(8); r.font.color.rgb = BLACK
    return r


def _field(p, instr):
    r = _ft_run(p, "")
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = instr
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    r._r.append(b); r._r.append(i); r._r.append(e)


def footer(d, txt, page_numbers=True):
    """Footer; mit Seitenzahl 'Seite X von Y' (Standard fuer LV)."""
    p = d.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if page_numbers:
        _ft_run(p, txt + " · Seite "); _field(p, "PAGE"); _ft_run(p, " von "); _field(p, "NUMPAGES")
    else:
        _ft_run(p, txt)


def jans_footer(d, doc_label, page_numbers=True):
    """Standard-JANS-Footer mit korrekter Adresse."""
    footer(d, f"{JANS_NAME} · {JANS_STREET} · {JANS_CITY} · {JANS_MAIL} · {JANS_WEB}"
           if not page_numbers else f"{JANS_NAME} · {doc_label}", page_numbers)


def build_anschreiben(path, *, empfaenger, betreff, anrede, absatz_einleitung,
                      grundlagen_satz, gleichwertig_bauseits, eingabefrist_satz,
                      beilagen, besichtigung=None, datum_ort="Zürich, 1. Juni 2026",
                      unterschrift="Raphael Jans"):
    """Baut ein Standard-Begleitschreiben mit allen Pflicht-Elementen.

    empfaenger: Liste Adresszeilen. beilagen: Liste (jede Beilage auf eigener Zeile).
    besichtigung: optionaler Satz zum Besichtigungstermin (oder None).
    """
    d = base_doc()
    para(d, JANS_NAME, bold=True)
    para(d, f"{JANS_STREET} · {JANS_CITY}", after=10)
    for line in empfaenger:
        para(d, line)
    para(d, "", after=2)
    para(d, datum_ort, align=WD_ALIGN_PARAGRAPH.RIGHT, after=10)
    para(d, betreff, bold=True, after=10)
    para(d, anrede, after=8)
    para(d, absatz_einleitung, after=6)
    para(d, grundlagen_satz, after=6)
    para(d, gleichwertig_bauseits, after=6)
    if besichtigung:
        para(d, besichtigung, after=6)
    para(d, eingabefrist_satz, after=10)
    para(d, "Freundliche Grüsse", after=8)
    para(d, unterschrift)
    para(d, JANS_NAME, after=10)
    para(d, "Beilagen:", size=10, before=2, after=2)
    for b in beilagen:
        para(d, "- " + b, size=10)
    jans_footer(d, "Begleitschreiben", page_numbers=False)
    d.save(path)


def build_lv(path, *, projekt, los_titel, bauherr_zeilen, positionen,
             einleitung=None, bauseits=None, datum_ort="Zürich, 1. Juni 2026",
             anbieter_zeilen=None):
    """Baut ein Leistungsverzeichnis (DOCX) mit leeren Preisspalten fuer den Anbieter.

    projekt:         z.B. "2620 Albertstrasse 7, 8008 Zürich"
    los_titel:       z.B. "LOS 273 — Schreinerarbeiten"
    bauherr_zeilen:  Liste Adresszeilen Bauherrschaft
    anbieter_zeilen: Liste Adresszeilen Anbieter — Stammzeile verlangt Objekt,
                     Bauherrschaft, Anbieter, Datum (10_dokumente-standard.md).
                     Optional (rueckwaertskompatibel): ohne Angabe entfaellt der Block.
    positionen:      Liste von Bereichen, je dict:
                       {"bereich": "...", "items": [
                           {"pos": "273.10", "bez": "...", "menge": "1", "einheit": "Stk",
                            "spez": "optional Spezifikation/Vermerk"}, ...]}
    einleitung:      optionaler Einleitungssatz. bauseits: Liste bauseitiger Leistungen.
    """
    d = base_doc()
    h1(d, "Leistungsverzeichnis")
    para(d, los_titel, bold=True, after=2)
    para(d, projekt, after=8)
    para(d, "Bauherrschaft:", bold=True, size=10, after=1)
    for z in bauherr_zeilen:
        para(d, z, size=10)
    if anbieter_zeilen:
        para(d, "Anbieter:", bold=True, size=10, before=6, after=1)
        for z in anbieter_zeilen:
            para(d, z, size=10)
    para(d, datum_ort, size=10, before=6, after=8)
    if einleitung:
        para(d, einleitung, after=8)
    para(d, "Die Einheits- und Gesamtpreise sind in den leeren Spalten einzutragen "
            "(exkl. MwSt). Mengen sind verbindlich, sofern nicht als 'ca.' bezeichnet.",
         size=10, after=8)

    # Spalten: Pos | Bezeichnung | Menge | Einh. | EP CHF | GP CHF  (Summe 170 mm)
    widths = [18, 84, 14, 14, 20, 20]
    aligns = ["l", "l", "r", "c", "r", "r"]
    for bereich in positionen:
        h2(d, bereich["bereich"])
        rows = [["Pos.", "Bezeichnung", "Menge", "Einh.", "EP CHF", "GP CHF"]]
        for it in bereich["items"]:
            bez = it["bez"]
            if it.get("spez"):
                bez = f"{bez}\n{it['spez']}"
            rows.append([it.get("pos", ""), bez, it.get("menge", ""),
                         it.get("einheit", ""), "", ""])
        table(d, rows, widths, header=True, aligns=aligns)

    # Summenblock
    para(d, "", before=4)
    sum_rows = [
        ["Total Positionen netto exkl. MwSt", ""],
        ["Rabatt / Skonto (falls gewährt)", ""],
        ["Zwischentotal exkl. MwSt", ""],
        [f"MwSt {MWST}", ""],
        ["Total inkl. MwSt", ""],
    ]
    table(d, sum_rows, [130, 40], header=False, aligns=["l", "r"])

    if bauseits:
        h2(d, "Bauseitige Leistungen (nicht im Auftrag enthalten)")
        for b in bauseits:
            para(d, "- " + b, size=10)

    para(d, "Gleichwertige Produkte sind zugelassen, sofern die technischen "
            "Spezifikationen vollumfänglich erfüllt werden.", size=10, before=10)
    jans_footer(d, los_titel)
    d.save(path)


def build_adressblatt(path, *, anbieter_zeilen):
    """Einseitiges Adressblatt fuer Fenstercouvert (pro Submittent). Adresse auf
    Hoehe des Sichtfensters positioniert (Rule adressblatt-submittent)."""
    d = base_doc()
    para(d, JANS_NAME, size=9)
    para(d, f"{JANS_STREET} · {JANS_CITY}", size=9, after=2)
    # Vertikaler Abstand bis Sichtfenster
    for _ in range(6):
        para(d, "", after=4)
    for z in anbieter_zeilen:
        para(d, z, size=12)
    d.save(path)


def build_antwortformular(path, *, projekt, los_titel, datum_ort="Zürich, 1. Juni 2026"):
    """Einheitliches Antwortformular, das jeder Anbieter ausfuellt — vereinfacht
    den spaeteren Offertenvergleich (gleiche Felder bei allen)."""
    d = base_doc()
    h1(d, "Antwortformular")
    para(d, los_titel, bold=True, after=2)
    para(d, projekt, after=8)
    para(d, "Bitte vollständig ausgefüllt mit dem bepreisten Leistungsverzeichnis "
            "retournieren.", size=10, after=8)
    felder = [
        ["Firma / Anbieter", ""],
        ["Kontaktperson", ""],
        ["Telefon / E-Mail", ""],
        ["Angebotssumme exkl. MwSt (CHF)", ""],
        [f"MwSt {MWST} (CHF)", ""],
        ["Angebotssumme inkl. MwSt (CHF)", ""],
        ["Rabatt / Skonto", ""],
        ["Verbindliche Lieferfrist / Termine", ""],
        ["Festpreis gültig bis (Preisbindung)", ""],
        ["Gewährleistung (Jahre)", ""],
        ["Zahlungskonditionen", ""],
        ["Referenzobjekte (vergleichbar)", ""],
        ["Vorbehalte / Bemerkungen", ""],
    ]
    table(d, felder, [70, 100], header=False, aligns=["l", "l"])
    para(d, "Ort / Datum:", size=10, before=14)
    para(d, "Rechtsverbindliche Unterschrift / Firmenstempel:", size=10, before=10)
    jans_footer(d, los_titel)
    d.save(path)


if __name__ == "__main__":
    print("jans_docx — Modul. Import in Build-Skripten verwenden. Siehe Docstring.")
