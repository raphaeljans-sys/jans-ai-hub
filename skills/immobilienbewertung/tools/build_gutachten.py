#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_gutachten.py — JANS Bewertungsgutachten (DOCX, WP-ueberlegene Optik, JANS-Layout)

Skill `immobilienbewertung`. Erzeugt aus einer JSON-Konfiguration ein vollstaendiges
Bewertungsgutachten im verbindlichen JANS-Layout (Cambria 11pt, A4, schwarz, rahmenlos)
mit Titelblatt (grosse Marktwert-Kennzahl) + frei zusammensetzbaren Bloecken.

Bewusst dem WP-Hedonic-Bericht ueberlegen: Augenschein/Qualitaetsprofil (1-5-Balken),
drei Verfahren, Marktwert-Synthese, Preisspanne, Bereinigungen, Glossar, Unterschrift.

Abhaengigkeit: python-docx (venv ~/.venvs/jansdocx). Layout-Bausteine aus dem Schwester-Skill
`ausschreibung` (jans_docx.py). PDF-Export anschliessend via scripts/docx2pdf.sh.

Aufruf:
  ~/.venvs/jansdocx/bin/python build_gutachten.py <config.json> <ausgabe.docx>

Block-Typen in config["blocks"] (Reihenfolge = Dokumentreihenfolge):
  {"h2": "1. Auftrag & Zweck"}                         Abschnitts-Ueberschrift
  {"para": "Text ..."}                                  Absatz (mehrere "para" hintereinander)
  {"para": "...", "bold": true}                          fetter Absatz (z.B. Hinweis)
  {"bullets": ["...", "..."]}                            Aufzaehlung (Bindestrich)
  {"stamm": [["Label","Wert"], ...]}                     Label-Wert-Tabelle (rahmenlos)
  {"table": {"rows": [[...],...], "widths":[..], "aligns":["l","r"], "header": true}}
  {"profile": [["Standard", 2.8, "Altbau, einfach"], ...]}   Qualitaetsprofil-Balken 1-5
  {"glossary": [["Marktwert","Definition ..."], ...]}    Glossar (Begriff fett + Text)

config-Kopf:
  titel, untertitel, objekt, marktwert, marktwert_label, status_hinweis (optional),
  stamm (Titelblatt-Stammdaten), doc_label (Footer), unterschrift (optional).
"""
import sys, os, json

_HERE = os.path.dirname(os.path.abspath(__file__))
# jans_docx (Layout-Bausteine) aus dem Schwester-Skill ausschreibung beziehen
sys.path.insert(0, os.path.join(_HERE, "..", "..", "ausschreibung", "tools"))
from jans_docx import (base_doc, para, h1, h2, table, jans_footer,  # noqa: E402
                       BLACK, FONT, JANS_NAME, JANS_STREET, JANS_CITY)
from docx.shared import Pt, Mm  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

CONTENT_MM = 170  # Inhaltsbreite bei A4 / 20 mm Raender

# ----------------------------------------------------------------- Bausteine

def _big_value(d, label, value):
    """Titelblatt-Kennzahl: kleines Label, darunter grosse Marktwert-Zahl (schwarz)."""
    para(d, label, size=12, bold=True, before=18, after=2)
    para(d, value, size=30, bold=True, after=8)


def _profile(d, rows, maxv=5):
    """Qualitaetsprofil als 1-5-Balken (gefuellt/leer), neutral schwarz.
    rows: Liste [name, wert(float), einstufung(str)]."""
    trows = [["Faktor (1–5)", "Profil", "Einstufung"]]
    for name, val, label in rows:
        try:
            v = float(val)
        except (TypeError, ValueError):
            v = 0.0
        filled = max(0, min(maxv, int(round(v))))
        bar = "■" * filled + "□" * (maxv - filled)
        wert = f"{bar}  {v:.1f}".rstrip("0").rstrip(".") if val not in (None, "") else bar
        trows.append([name, f"{bar}  {val}", label])
    table(d, trows, [46, 64, 60], header=True, aligns=["l", "l", "l"])


def _glossary(d, items):
    for begriff, text in items:
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{begriff}: ")
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(10)
        r.font.color.rgb = BLACK
        r2 = p.add_run(text)
        r2.font.name = FONT
        r2.font.size = Pt(10)
        r2.font.color.rgb = BLACK


def _bullets(d, items):
    for it in items:
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Mm(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"-  {it}")
        r.font.name = FONT
        r.font.size = Pt(11)
        r.font.color.rgb = BLACK


# ----------------------------------------------------------------- Hauptbau

def build(cfg, out_path):
    d = base_doc()

    # ---- Titelblatt
    para(d, JANS_NAME, size=11, bold=True, after=0)
    para(d, f"{JANS_STREET} · {JANS_CITY}", size=10, after=10)
    h1(d, cfg.get("titel", "Bewertungsgutachten"))
    if cfg.get("untertitel"):
        para(d, cfg["untertitel"], size=12, bold=True, after=2)
    if cfg.get("objekt"):
        para(d, cfg["objekt"], size=12, after=4)
    if cfg.get("marktwert"):
        _big_value(d, cfg.get("marktwert_label", "Marktwert"), cfg["marktwert"])
    if cfg.get("status_hinweis"):
        para(d, cfg["status_hinweis"], size=10, bold=True, before=6, after=6)
    if cfg.get("stamm"):
        table(d, cfg["stamm"], [46, 124], header=False, aligns=["l", "l"])

    # ---- Bloecke
    for blk in cfg.get("blocks", []):
        if "h2" in blk:
            h2(d, blk["h2"])
        elif "para" in blk:
            para(d, blk["para"], size=11, bold=blk.get("bold", False), after=4)
        elif "bullets" in blk:
            _bullets(d, blk["bullets"])
        elif "stamm" in blk:
            table(d, blk["stamm"], blk.get("widths", [46, 124]), header=False,
                  aligns=blk.get("aligns", ["l", "l"]))
        elif "table" in blk:
            t = blk["table"]
            table(d, t["rows"], t["widths"], header=t.get("header", True),
                  aligns=t.get("aligns"))
        elif "profile" in blk:
            _profile(d, blk["profile"])
        elif "glossary" in blk:
            _glossary(d, blk["glossary"])

    # ---- Unterschrift
    if cfg.get("unterschrift"):
        para(d, "", after=10)
        for line in cfg["unterschrift"]:
            para(d, line, size=11, after=0)

    jans_footer(d, cfg.get("doc_label", "Bewertungsgutachten"))
    d.save(out_path)
    print(f"OK {out_path}")


def main():
    if len(sys.argv) != 3:
        print("Aufruf: build_gutachten.py <config.json> <ausgabe.docx>", file=sys.stderr)
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        cfg = json.load(f)
    build(cfg, sys.argv[2])


if __name__ == "__main__":
    main()
